#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
预算调整 docx 生成器(填充模板)。

用法:脚本从 stdin 读取一个 JSON,填充 template/预算调整-template.docx,
把生成的 docx 二进制写到 stdout。失败时写错误信息到 stderr 并以非 0 退出。

输入 JSON 结构:
{
  "templatePath": "template/预算调整-template.docx",
  "title": "总预算调整表",            // 文档首段标题
  "project": {
    "name": "项目名称",
    "projectType": "项目类型",
    "undertakingUnit": "承担单位",
    "ownerName": "项目负责人",
    "researchPeriod": "2026.01-2028.12",
    "totalFundWan": "10.00",          // 项目总经费(万元)
    "annualFundWan": "10.00"          // 年度预算经费(万元)
  },
  "reason": "调整原因说明文字",
  "rows": [                           // 调整明细,每条对应模板一个数据行
    {
      "subjectTitle": "设备费",        // 预算科目(二级标题)
      "productName": "设备购置费",     // 品名(叶科目名;二级无子则空串)
      "originWan": "6.00",            // 原预算金额(万元)
      "adjustedWan": "5.00",          // 调整后金额(万元)
      "adjustWan": "-1.00"            // 调整金额(万元,调整后-原)
    }
  ],
  "totalAdjustWan": "0.00"            // 合计调整金额(万元)
}

金额均为"万元、两位小数"字符串,由 Node 端算好传入。
"""
import json
import os
import sys

from docx import Document
from docx.oxml.ns import qn


def _set_cell_text(cell, text):
    """设置单元格文本:保留首个 paragraph 的格式,清空多余 run,写入文本。

    python-docx 的 cell.text = ... 会重建段落丢失字体格式,故手动改 run 文本。
    """
    text = text if text is not None else ''
    para = cell.paragraphs[0]
    # 清空该段所有 run,只保留首个 run 的 rPr 作为格式模板。
    runs = para.runs
    if runs:
        first = runs[0]
        # 删除首个 run 之后的所有 run。
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
        first.text = text
    else:
        # 无 run(空段),新建一个。
        para.add_run(text)


def _clone_row(template_row, table):
    """深拷贝一个 tr(含格式),插入到该行之后,返回新 row 对象。"""
    from copy import deepcopy

    new_tr = deepcopy(template_row._tr)
    template_row._tr.addnext(new_tr)
    # 用新 tr 在父节点的位置定位返回对应 row。
    parent = template_row._tr.getparent()
    # 重新构建 rows:python-docx 的 table.rows 会按 tr 顺序重新生成。
    return table.rows[parent.index(new_tr)]


def _remove_row(row):
    """从表格移除一行。"""
    tr = row._tr
    tr.getparent().remove(tr)


def fill_template(data):
    template_path = data.get('templatePath') or 'template/预算调整-template.docx'
    doc = Document(template_path)
    table = doc.tables[0]

    project = data.get('project', {})
    rows = data.get('rows', [])

    # ---- 文档首段标题(模板首段为空或占位,替换为维度标题) ----
    title = data.get('title', '')
    if title and doc.paragraphs:
        first = doc.paragraphs[0]
        if first.runs:
            first.runs[0].text = title
            for r in first.runs[1:]:
                r._r.getparent().remove(r._r)
        else:
            first.add_run(title)

    # ---- 项目情况区(logical 列布局,见脚本顶部注释) ----
    # 行0: c1=项目名称(标签) → 值填 c3
    _set_cell_text(table.cell(0, 3), project.get('name', ''))
    # 行1: c1=项目类型 → c3 ; c6=承担单位 → c8
    _set_cell_text(table.cell(1, 3), project.get('projectType', ''))
    _set_cell_text(table.cell(1, 8), project.get('undertakingUnit', ''))
    # 行2: c1=项目负责人 → c3 ; c6=研究周期 → c8
    _set_cell_text(table.cell(2, 3), project.get('ownerName', ''))
    _set_cell_text(table.cell(2, 8), project.get('researchPeriod', ''))
    # 行3: c1=项目总经费 → c3 ; c6=年度预算经费 → c8
    _set_cell_text(table.cell(3, 3), project.get('totalFundWan', ''))
    _set_cell_text(table.cell(3, 8), project.get('annualFundWan', ''))

    # ---- 调整内容表:数据行模板在行7,空行在 7..19(共13行) ----
    DATA_START = 7  # 第一个数据行索引
    BLANK_ROWS = 13  # 模板预留空数据行数(行7-19)

    # 需要的总数据行数 = 明细数(至少 1,保证表格结构)。
    need = max(len(rows), 1)

    # 调整数据行数量:行7 是模板行,克隆 need-1 份插到其后;若 need<13 删除多余空行。
    template_data_row = table.rows[DATA_START]
    # 先克隆到 need 行(在行7 之后依次插入)。
    for i in range(1, need):
        _clone_row(table.rows[DATA_START + i - 1], table)
    # 此时 DATA_START..DATA_START+need-1 是数据行;其后原本的空行(直到 DATA_START+13-1)需删除多余。
    # 删除 need..BLANK_ROWS-1 之间的多余空行(它们现在索引为 DATA_START+need .. DATA_START+BLANK_ROWS-1)。
    extra = BLANK_ROWS - need
    for _ in range(extra):
        # 每次删 DATA_START+need 这一行(删除后下一行补位)。
        _remove_row(table.rows[DATA_START + need])

    # 填充每行数据(logical 列: 0序号 1预算科目 2原品名 4原金额 5调后品名 8调后金额 9调整金额)。
    for i, row in enumerate(rows):
        r = DATA_START + i
        _set_cell_text(table.cell(r, 0), str(i + 1))
        _set_cell_text(table.cell(r, 1), row.get('subjectTitle', ''))
        _set_cell_text(table.cell(r, 2), row.get('productName', ''))
        _set_cell_text(table.cell(r, 4), row.get('originWan', ''))
        _set_cell_text(table.cell(r, 5), row.get('productName', ''))  # 调后品名同原品名
        _set_cell_text(table.cell(r, 8), row.get('adjustedWan', ''))
        _set_cell_text(table.cell(r, 9), row.get('adjustWan', ''))

    # ---- 合计行:原本在行20,前面删/增行后位置变化,按标签定位 ----
    # 找"合计金额"所在行,在其调整金额列(log9)填合计。
    for ri, row in enumerate(table.rows):
        if row.cells[0].text.strip().startswith('合计金额'):
            _set_cell_text(table.cell(ri, 9), data.get('totalAdjustWan', ''))
            break

    # ---- 调整原因:定位模板里占位为"XX。"的单元格替换(精确匹配,避免误伤含 XX 的其它值)。
    reason = data.get('reason', '')
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            if cell.text.strip() == 'XX。' or cell.text.strip() == 'XX':
                _set_cell_text(cell, reason)
                break

    # 输出到 stdout(二进制)。
    out = sys.stdout.buffer
    doc.save(out)


def main():
    raw = sys.stdin.read()
    if not raw.strip():
        sys.stderr.write('未收到输入 JSON\n')
        sys.exit(2)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        sys.stderr.write(f'JSON 解析失败: {e}\n')
        sys.exit(2)
    try:
        fill_template(data)
    except Exception as e:
        sys.stderr.write(f'生成 docx 失败: {e}\n')
        raise


if __name__ == '__main__':
    main()
